import re
import io
import pandas as pd
import streamlit as st
from openpyxl.styles import Font
import os 
from docx import Document
import docx
from openpyxl import Workbook


# ---------------------------------------------------------
# Backend Processing Functions
# ---------------------------------------------------------

def read_docx_file(file):
    """Extract text from a DOCX file-like object."""
    doc = docx.Document(file)
    return "\n".join(para.text for para in doc.paragraphs)

def get_corresponding_bank(preamble):
    """
    Build the corresponding bank from preamble lines.
    Drops any lines starting with "20:".
    Expects the first remaining line to be the bank name and the second to be the currency descriptor.
    """
    filtered = [line for line in preamble if not line.startswith("20:")]
    if len(filtered) >= 2:
        bank_name = f"{filtered[0]} - {filtered[1]}"
        return bank_name, filtered[1]
    elif len(filtered) == 1:
        return filtered[0], None
    else:
        return "Unknown Bank", None

def clean_sheet_name(sheet_name):
    """
    Clean up the bank name so it is a valid Excel sheet name:
      - Replace invalid characters with underscores.
      - Limit length to 31 characters (Excel's limit).
    """
    cleaned = re.sub(r'[\\/*?:\[\]]', '_', sheet_name)
    return cleaned[:31]

def process_message(block):
    """
    Processes a single SWIFT message block.
    Returns a dict with:
      - bank_name
      - opening_balance (from :60F:)
      - closing_balance (from :62F:)
      - transactions: a list of dicts containing transaction data (from :61: tags)
    """
    preamble = []
    swift_lines = []
    for line in block:
        if line.startswith(":"):
            swift_lines.append(line)
        else:
            if not swift_lines:
                preamble.append(line)
            else:
                swift_lines.append(line)
    
    # Skip block if none of the relevant SWIFT tags are present.
    if not any(tag in " ".join(swift_lines) for tag in [":60F:", ":61:", ":62F:"]):
        return None

    corresponding_bank, _ = get_corresponding_bank(preamble)
    opening_balance = None
    closing_balance = None
    currency_code = None
    transactions = []

    # Regex for :61: tag with optional 4-digit entry date.
    tx_pattern = r"^:61:(\d{6})(\d{4})?((?:C(?:[CDPR])?)|(?:D(?:[RPD])?))([\d,]+)(.+)$"
    
    i = 0
    while i < len(swift_lines):
        line = swift_lines[i]
        if line.startswith(":60F:"):
            m = re.match(r":60F:[CD]{1,2}(\d{6})([A-Z]{3,})([\d,]+)", line)
            if m:
                currency_code = m.group(2)
                amount_str = m.group(3).replace(',', '.')
                opening_balance = float(amount_str)
        elif line.startswith(":62F:"):
            m = re.match(r":62F:[CD]{1,2}(\d{6})([A-Z]{3,})([\d,]+)", line)
            if m:
                amount_str = m.group(3).replace(',', '.')
                closing_balance = float(amount_str)
        elif line.startswith(":61:"):
            m_tx = re.match(tx_pattern, line)
            if m_tx:
                tx_date = m_tx.group(1)
                # m_tx.group(2) is the optional entry date (unused here)
                tx_code = m_tx.group(3).strip().upper()
                tx_amount = float(m_tx.group(4).replace(',', '.'))
                reference = m_tx.group(5).strip()
                ordering_customer = ""
                if (i + 1) < len(swift_lines) and not swift_lines[i+1].startswith(":"):
                    ordering_customer = swift_lines[i+1].strip()
                    i += 1

                # Explicitly handle known codes.
                if tx_code in {"D", "DR", "DP", "DD"}:
                    debit_amt = tx_amount
                    credit_amt = None
                elif tx_code in {"C", "CC", "CD", "CP", "CR"}:
                    debit_amt = None
                    credit_amt = tx_amount
                else:
                    if tx_code.startswith("D"):
                        debit_amt = tx_amount
                        credit_amt = None
                    else:
                        debit_amt = None
                        credit_amt = tx_amount

                transactions.append({
                    "Date": tx_date,
                    "Currency": currency_code if currency_code else "",
                    "Ordering Customer": ordering_customer,
                    "Swift_Credit": credit_amt,
                    "Swift_Debit": debit_amt,
                    "Reference": reference
                })
        i += 1

    return {
        "bank_name": corresponding_bank,
        "opening_balance": opening_balance,
        "closing_balance": closing_balance,
        "transactions": transactions
    }

def process_swift_message(swift_message):
    """
    Process the entire SWIFT text.
    Splits the text into blocks, processes each block, and aggregates data by corresponding bank.
    """
    swift_message = re.sub(r"(:\d{2}[A-Z]{0,3}:)", r"\n\1", swift_message).strip()
    all_lines = [line.strip() for line in swift_message.splitlines() if line.strip()]
    messages = []
    current_block = []
    for line in all_lines:
        if not current_block:
            current_block.append(line)
        else:
            if (not line.startswith(":")) and any(l.startswith(":62F:") for l in current_block):
                messages.append(current_block)
                current_block = [line]
            else:
                current_block.append(line)
    if current_block:
        messages.append(current_block)
    
    bank_data = {}
    for block in messages:
        result = process_message(block)
        if result is not None:
            bname = result["bank_name"]
            bank_data.setdefault(bname, []).append(result)
    return bank_data

# ---------------------------------------------------------
# Streamlit Frontend
# ---------------------------------------------------------

st.set_page_config(page_title="SWIFT Extractor App", layout="centered")


st.markdown(
    """
    <style>
    footer, footer a {
        color: #999 !important;
        pointer-events: none !important;
        user-select: none !important;
        text-decoration: none !important;
        cursor: default !important;
    }
    </style>
    """,
    unsafe_allow_html=True
)

st.markdown("<h1 style='text-align: center;'>SWIFT Extractor App</h1>", unsafe_allow_html=True)



# Sidebar navigation 
with st.sidebar:
    st.markdown(
        """
        <div style="display: flex; flex-direction: column; align-items: center; text-align: center; padding-top: 20px;">
            <h2 style="margin-bottom: 10px;">ℹ️ About The SWIFT Extractor App</h2>
            <p style="max-width: 250px;">
                <strong>The SWIFT Extractor App</strong> automates the extraction, formatting, and export of swift statement data to Excel for reconciliation and financial review.
            </p>
            <ul style="list-style: none; padding: 0; margin-top: 10px;">
                <li>📄 Parse balances and transactions</li>
                <li>📥 Download Word/Excel templates </li>
                <li>📊 Export data for reconciliation</li>
            </ul>
        </div>
        """, unsafe_allow_html=True
    )


# ————— Word Template —————
def generate_word_template():
    doc = Document()
    doc.add_heading("SWIFT Transactions Template", 0)
    doc.add_paragraph(
               "Please use the format below to paste your SWIFT transactions.\n\n"
                "BOG \nUSD\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"

                "BOG \nGBP\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"

                "BOG \nEUR\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"

                "BOG \nYUAN\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"

                "ABSA \nUSD\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"

                "JP MORGAN MAIN \nUSD\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"

                "JP MORGAN SUB \nUSD\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"

                "GIB \nUSD\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"

                "GIB \nEUR\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"

                "STANDARD BANK SA \nUSD\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"

                "STANDARD CHARTERED UK SCBLGB2L \nUSD\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"

                "STANDARD CHARTERED UK \nGBP\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"

                "STANDARD CHARTERED UK \nYUAN\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"
                
                "STANDARD CHARTERED UK \nSCBL\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"

                "STANDARD CHARTERED UK \nEUR\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"

                "GTB UK GTBIGB2L \nUSD\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"

                "GTB UK GTBIGB2L SUB \nUSD\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"

                "GTB UK GTBIGB2L SUB 1 \nUSD\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"

                "GTB UK SUB 2 \nUSD\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"

                "GTB UK \nEUR\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"

                "GTB UK \nGBP\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"

                "GTB UK SUB \nGBP\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"

                "GTB UK SUB \nEUR\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"

                "CITI LDN \nUSD\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"

                "CITI LDN \nGBP\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"

                "CITI LDN \nEUR\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"

                "CITI LDN \nZAR\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"

                "CITI LDN FX \nUSD\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"

                "CITI LDN FX \nGBP\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"

                "CITI LDN FX \nEUR\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"

                "RMB \nUSD\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"

                "CITI NY \nMASTERCARD\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"

                "CITI NY \nSETTLEMENT\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"

                "CITI NY \nMAIN\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"

                "CITI NY SUB \nUSD\n:20:REFERENCE123\n:60F:C230501USD1234,56\n"
                ":61:230502C123,45REFERENCE\nCustomer Name\n:62F:C230530USD5678,90\n\n"
    )
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ————— Excel Template —————
def generate_excel_template():
    bank_names = [
                "BOG - USD", "BOG - GBP", "BOG - EUR", "BOG - YUAN", "ABSA - USD", "JP MORGAN MAIN - USD",
                "JP MORGAN SUB - USD", "GIB - USD", "GIB - EUR", "STANDARD BANK SA - USD", 
                "STANDARD CHARTERED UK SCBLGB2L", "STANDARD CHARTERED UK - GBP",
                "STANDARD CHARTERED UK - SCBLUS3", "STANDARD CHARTERED UK - EUR", "GTB UK - USD SUB 1",
                "GTB UK GTBIGB2L SUB - USD", "GTB UK GTBIGB2L SUB 1 - GBP", "GTB UK GTBIGB2L SUB 2 - USD",
                "GTB UK - EUR", "GTB UK - GBP", "GTB UK SUB - EUR", "CITI LDN - USD", "CITI LDN - GBP", "CITI LDN - EUR",
                "CITI LDN - ZAR", "CITI LDN FX - USD", "CITI LDN FX - GBP",
                "CITI LDN FX - EUR", "RMB - USD", "CITI NY - MASTERCARD",
                "CITI NY - SETTLEMENT", "CITI NY - MAIN", "CITI NY SUB - USD"
    ]
    wb = Workbook()
    # remove default sheet
    wb.remove(wb.active)
    for bank in bank_names:
        wb.create_sheet(title=bank)
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer

# ————— UI: Two‑Column Downloads —————
with st.expander("Download Templates", expanded=False):
    col1, col2 = st.columns(2)

    # Word
    with col1:
        st.markdown("#### Word Template")
        word_buf = generate_word_template()
        st.download_button(
            label="📄 Download .docx",
            data=word_buf,
            file_name="swift_template.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # Excel
    with col2:
        st.markdown("#### Excel Template")
        excel_buf = generate_excel_template()
        st.download_button(
            label="📑 Download .xlsx",
            data=excel_buf,
            file_name="corresponding_banks_template.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )


uploaded_file = st.file_uploader("Upload a DOCX file", type=["docx"])

if uploaded_file is not None:
    with st.spinner("Processing file..."):
        swift_text = read_docx_file(uploaded_file)
        bank_data = process_swift_message(swift_text)
    
    st.success("File processed successfully!")
    
    if bank_data:
        with st.expander("Processed Banks", expanded=False):
            bank_names = list(bank_data.keys())
            st.write(bank_names)
        
        selected_bank = st.selectbox("Select a bank to preview", list(bank_data.keys()))
        if selected_bank in bank_data:
            block = bank_data[selected_bank][0]
            
            st.subheader(f"Opening & Closing Balance for {selected_bank}")
            balance_df = pd.DataFrame({
                "Balance": ["Opening Balance", "Closing Balance"],
                "Value": [
                    block["opening_balance"] if block["opening_balance"] is not None else "",
                    block["closing_balance"] if block["closing_balance"] is not None else ""
                ]
            })
            st.dataframe(balance_df)
            
            st.subheader(f"Transactions for {selected_bank}")
            tx_data = block["transactions"]
            if tx_data:
                tx_df = pd.DataFrame(tx_data)
            else:
                tx_df = pd.DataFrame(columns=["Date", "Currency", "Ordering Customer", "Swift_Credit", "Swift_Debit", "Reference"])
            
            # Ensure the columns exist and fill missing values.
            if "Ordering Customer" not in tx_df.columns:
                tx_df["Ordering Customer"] = ""
            if "Reference" not in tx_df.columns:
                tx_df["Reference"] = ""
            tx_df["Ordering Customer"] = tx_df["Ordering Customer"].fillna("").astype(str).str.strip()
            tx_df["Reference"] = tx_df["Reference"].fillna("").astype(str).str.strip()
            tx_df["Narration"] = (tx_df["Ordering Customer"] + " " + tx_df["Reference"]).str.strip()
            st.dataframe(tx_df)
        
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine="openpyxl") as writer:
            for bank, blocks in bank_data.items():
                sheet_name = clean_sheet_name(bank)
                start_row = 0
                for blk in blocks:
                    obal = blk["opening_balance"]
                    cbal = blk["closing_balance"]
                    txs = blk["transactions"]
                    balance_df = pd.DataFrame({
                        "Balance": ["Opening Balance", "Closing Balance"],
                        "Value": [obal if obal is not None else "", 
                                  cbal if cbal is not None else ""]
                    })
                    balance_df.to_excel(writer, sheet_name=sheet_name, startrow=start_row, index=False)
                    start_row += len(balance_df) + 1
                    if txs:
                        tx_df = pd.DataFrame(txs, columns=["Date", "Currency", "Ordering Customer", "Swift_Credit", "Swift_Debit", "Reference"])
                    else:
                        tx_df = pd.DataFrame(columns=["Date", "Currency", "Ordering Customer", "Swift_Credit", "Swift_Debit", "Reference"])
                    if "Ordering Customer" not in tx_df.columns:
                        tx_df["Ordering Customer"] = ""
                    if "Reference" not in tx_df.columns:
                        tx_df["Reference"] = ""
                    tx_df["Ordering Customer"] = tx_df["Ordering Customer"].fillna("").astype(str).str.strip()
                    tx_df["Reference"] = tx_df["Reference"].fillna("").astype(str).str.strip()
                    tx_df["Narration"] = (tx_df["Ordering Customer"] + " " + tx_df["Reference"]).str.strip()
                    tx_df.to_excel(writer, sheet_name=sheet_name, startrow=start_row, index=False)
                    start_row += len(tx_df) + 2

            for sheet in writer.book.worksheets:
                for row in sheet.iter_rows():
                    for cell in row:
                        cell.font = Font(name="Arial", size=10)
        output.seek(0)

        # Create dynamic download filename 
        original_filename = os.path.splitext(uploaded_file.name)[0]
        processed_filename = f"{original_filename}_processed.xlsx"

        # Download button with dyanmic name 
        st.download_button(
            "Download Excel file",
            data=output,
            file_name=processed_filename,
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
    else:
        st.warning("No SWIFT transactions found in the file.")